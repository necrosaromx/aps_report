#!/usr/bin/python
# -*- coding: utf-8 -*-
#Autor: Saul Vargas Leon
from pdf2image import convert_from_path, convert_from_bytes
from docx import Document
from docx.shared import Inches, Pt
from os import listdir, rmdir, remove 
from os.path import isfile, join
import logging
from tqdm import tqdm


def generadoc(configuracion, archivosalida, grupos, fechas):
    """
        Genera el documento completo 
    """
    documento = Document(configuracion['plantilla'])
    fuente = documento.styles['Normal'].font
    fuente.name = "Calibri"
    fuente.size = Pt(36)
    documento.add_paragraph('Reporte Diario de Monitoreo PRAVAIL').alignment = 1
    documento.add_paragraph(configuracion['hostname']).alignment = 1
    documento.add_paragraph('')
    documento.add_paragraph('')
    documento.add_paragraph('Periodo:').alignment = 1
    documento.add_paragraph(str(fechas['dia_ini']) 
                            + '/' 
                            + str(fechas['mes_ini']) 
                            + '/' 
                            + str(fechas['year_ini']) 
                            + '  ' 
                            + str(fechas['marcial_ini']) 
                            + ':00 HORAS').alignment = 1
    documento.add_paragraph(str(fechas['dia_fin']) 
                            + '/' 
                            + str(fechas['mes_fin']) 
                            + '/' 
                            + str(fechas['year_fin']) 
                            + '  ' 
                            + str(fechas['marcial_fin']) 
                            + ':00 HORAS').alignment = 1
    for g in tqdm(grupos, desc="Generando Doc", unit="gp"):
        logging.debug("Agregando %s", g[0])
        documento.add_heading(g[0])
        pdf = join(configuracion['output_dir'], g[0] + '.pdf')
        images = convert_from_path(pdf, fmt="png")
        for i, pagina in enumerate(images):
            img = join(configuracion['output_dir'], g[0] + str(i + 1) + '.png')
            pagina.save(img)
            logging.debug("Agregando Grupo: %s/ Imagen: %s al reporte Word.", g[0], img)
            documento.add_picture(img, width=Inches(5.77)) 
            logging.debug("Borrando archivo imagen")
            remove(img)
        if not configuracion['guardapdf']:
            remove(pdf)
    documento.save(join(configuracion['output_dir'], archivosalida))
    
